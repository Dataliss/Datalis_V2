from base_agent import BaseAgent
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64

class AuditorAgent(BaseAgent):
    """Auditor Agent implementation"""
    
    @property
    def name(self):
        return "Auditor Agent"
    
    @property
    def system_prompt(self):
        return """You are an expert auditor. Provide professional, accurate insights, with citations and references."""
    
    def determine_audit_framework(self, document_texts, audit_type=None):
        """Identify the most appropriate audit framework based on document content."""
        # Combine texts and get a representative sample
        combined_text = "\n\n".join([f"Document: {text[:1500]}..." for text in document_texts])
            
        prompt = (
            f"Based on these financial document excerpts:\n\n{combined_text}\n\n"
            f"Determine the most appropriate audit framework for {audit_type or 'financial audit'}. "
            f"Examples include SA 700, AS 1, Ind AS 109, etc. "
            f"Provide the framework name and a brief explanation of why it's appropriate."
        )
            
        return self.llm_service.get_response(prompt, self.system_prompt)
    
    def generate_audit_report_docx(self, audit_type, document_texts, framework, company_info=None):
        """Generate a professional DOCX audit report."""
        doc = Document()
            
        # Add title
        title = doc.add_heading('Audit Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Add subtitle with audit type
        subtitle = doc.add_heading(f'{audit_type} Assessment', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Add date
        from datetime import datetime
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
        
        # Add company information if available
        if company_info and not company_info.get("skipped", False):
            # Company info section implementation...
            doc.add_heading('Company Information', 1)
            # ... (rest of company info implementation)
        
        # Add executive summary
        doc.add_heading('Executive Summary', 1)
        summary_text = self.llm_service.get_response(
            f"Create an executive summary for an {audit_type} report based on these documents:\n\n" +
            "\n\n".join([text[:1500] + "..." for text in document_texts]) +
            "\n\nWrite a professional, concise executive summary (3-4 paragraphs).",
            self.system_prompt
        )
        doc.add_paragraph(summary_text)
            
        # Add scope section
        doc.add_heading('Scope of Audit', 1)
        scope_text = self.llm_service.get_response(
            f"Create a scope section for an {audit_type} using framework {framework}. "
            "Describe what was covered in the audit, methodology used, and time period.",
            self.system_prompt
        )
        doc.add_paragraph(scope_text)
            
        # Add findings section
        doc.add_heading('Key Findings', 1)
        findings_text = self.llm_service.get_response(
            f"Generate key findings for an {audit_type} based on these documents:\n\n" +
            "\n\n".join([text[:1500] + "..." for text in document_texts]) +
            "\n\nCreate 3-5 significant findings with details."
            "\n\nProvide specific citations or references to the documents where applicable.",
            self.system_prompt
        )
        doc.add_paragraph(findings_text)
            
        # Add recommendations
        doc.add_heading('Recommendations', 1)
        recommendations_text = self.llm_service.get_response(
            f"Based on an {audit_type} audit with these findings:\n\n{findings_text}\n\n"
            "Provide 3-5 specific, actionable recommendations.",
            self.system_prompt
        )
        doc.add_paragraph(recommendations_text)
            
        # Add conclusion
        doc.add_heading('Conclusion', 1)
        conclusion_text = self.llm_service.get_response(
            f"Write a conclusion for an {audit_type} audit report that summarizes the overall assessment, "
            f"significance of findings, and next steps. Keep it professional and concise."
            f"Add final thoughts on the audit process and references to the documents reviewed.",
            self.system_prompt
        )
        doc.add_paragraph(conclusion_text)
            
        # Add digital signature if available
        if company_info and company_info.get("digital_signature"):
            doc.add_heading('Auditor Signature', 1)
            signature_paragraph = doc.add_paragraph()
            signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
            try:
                # Decode the base64 signature
                signature_data = base64.b64decode(company_info["digital_signature"])
                        
                # Create a temporary file for the signature
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                    temp_file.write(signature_data)
                    temp_file_path = temp_file.name
                            
                # Add the signature image to the document
                signature_paragraph.add_run().add_picture(temp_file_path, width=Inches(2))
                            
                # Clean up the temporary file
                os.unlink(temp_file_path)
                            
                # Add CA name and ID below signature
                if company_info.get("ca_name"):
                    ca_info = doc.add_paragraph()
                    ca_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ca_info.add_run(company_info["ca_name"])
                                    
                    if company_info.get("ca_id"):
                        ca_info.add_run(f"\nCA Membership: {company_info['ca_id']}")
                                    
                    if company_info.get("ca_firm"):
                        ca_info.add_run(f"\n{company_info['ca_firm']}")
            except Exception as e:
                print(f"Error adding signature to document: {str(e)}")
            
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file_path = temp_file.name
        temp_file.close()
        doc.save(temp_file_path)
            
        return temp_file_path
    
    def analyze_documents(self, document_texts, audit_type=None):
        """Analyze document content for audit insights."""
        # Combine texts and get a representative sample
        combined_text = "\n\n".join([text[:3000] for text in document_texts])
        
        prompt = (
            f"Analyze these financial documents for a {audit_type or 'Financial Statement Audit'}:\n\n{combined_text}\n\n"
            "Provide key observations, potential risks, and compliance issues."
        )
                
        return self.llm_service.get_response(prompt, self.system_prompt)
    
    def generate_suggested_questions(self, text):
        """Generate relevant audit questions based on document content."""
        # Limit text length for API
        sample_text = text[:4000]
            
        prompt = (
            f"Based on the following financial information:\n\n{sample_text}\n\n"
            "Generate 5 key questions an auditor should ask. "
            "Format each question as a numbered list (1., 2., etc.). "
            "Focus on potential risk areas, compliance concerns, and areas needing clarification."
        )
            
        return self.llm_service.get_response(prompt, self.system_prompt)
